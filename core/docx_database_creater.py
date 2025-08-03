from docx_table_reader import get_document
from docx_table_reader import extract_table_info
from docx.oxml.text.paragraph import CT_P
from sentence_transformers import SentenceTransformer
import os
import re
import faiss
import numpy as np

#checking text validation
def gibbrish_detector(text) -> bool:
    #contain special character or number --> remove
    if re.fullmatch(r"[^\w\sÀ-Ỵà-ỵ]{3,}", text):
        return True
    if len(re.findall(r"[A-Z]{2,}", text)) > 5:
        return True
    if re.search(r"[\^_~@#\$%]", text):
        return True
    return False

def is_valid_text(text) -> bool:
    if not text.strip():
        return False
    if len(text.strip()) < 5:
        return False
    if gibbrish_detector(text):
        return False
    if re.search(r"[a-zA-ZÀ-Ỵà-ỵ]", text):  # có ký tự chữ
        return True
    return False

#find the header of table
def get_table_header(document, table_index, max_lookback=5) -> str:
    #find the indicated table
    table = document.tables[table_index]
    table_element = table._element

    #scan all element in the document
    body_elements = list(document.element.body.iterchildren())

    #find the position of indicated table in XML tree
    table_position = None
    for i, el in enumerate(body_elements):
        if el == table_element:
            table_position = i
            break

    #tray back to find the header of table
    if table_position is not None:
        lookback = 0
        for i in range(table_position - 1, -1, -1):
            #only tray back in a specific range
            if lookback > max_lookback:
                break
            current_el = body_elements[i]
            #in case of this element is a paragraph
            if isinstance(current_el, CT_P):
                para_candidate = current_el.text.strip()
                #checking it text validation
                if is_valid_text(para_candidate):
                    #it might be the header of the table
                    return para_candidate
                lookback += 1

    return f"Table {table_index}"

#convert table to text for easier embedding
def table_to_text(table_data, title="") -> str:
    # table_data: List[List[str]]
    flat_lines = [", ".join(row) for row in table_data]
    table_text = f"{title}\n" + "\n".join(flat_lines)
    return table_text


#create vector embedding for input file
def create_embedding_vector(input_document, input_model='all-MiniLM-L6-v2') -> tuple:
    model = SentenceTransformer(input_model)             #the model, change it if you want
    embeddings = []
    metadatas = []

    #for docx file
    for idx, table in enumerate(input_document.tables):
        data = extract_table_info(tables=input_document.tables, table_index=idx)
        title = get_table_header(document=input_document, table_index=idx, max_lookback=5)
        text = table_to_text(table_data=data, title=title)

        vector = model.encode(text)
        embeddings.append(vector)
        metadatas.append({
            "table_index": idx,
            "title": title,
            "text": text
        })

    return embeddings, metadatas

#todo: create embedding vector database
def get_file_path():
    #read input path
    try:
        base_dir = os.path.dirname(__file__)
        input_path = os.path.abspath(os.path.join(base_dir, '..', '000000014601738_VI_BaoCaoTaiChinh_KiemToan_2024_HopNhat_14032025110908.docx'))
        print('file located')
        return input_path

    except Exception as e:
        print('cant locate file')

def create_database(input_path):
    #read document
    input_path = input_path
    try:
        document = get_document(path=input_path)
        print('document reading successful')

    except Exception as e:
        print('cant read document')

    #embedding
    try:
        input_model = 'all-MiniLM-L6-v2'
        embeddings, metadatas = create_embedding_vector(input_document=document, input_model=input_model)
        print('embedding vector successful')

    except Exception as e:
        print('cant create embedding vector')
        print(f'Error: {e}')

    try:
        # Convert to numpy
        embedding_matrix = np.stack(embeddings).astype("float32")
        faiss.normalize_L2(embedding_matrix)

        # FAISS index
        dimension = embedding_matrix.shape[1]
        index = faiss.IndexFlatIP(dimension)
        index.add(embedding_matrix)
        print('database created')
        return index, metadatas

    except Exception as e:
        print('cant create database')
        print(f'Error: {e}')