from docx import Document
import pandas as pd
import os
from data_process.creat_vecdb import add_new_doc
import json
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
import fitz
import pdfplumber
import fitz 
from sentence_transformers import SentenceTransformer, util
import re
model_embedding = SentenceTransformer('all-MiniLM-L6-v2')
pdf_data_path = "data"
vector_db_path = "vectorstores/db_faiss"


def read_docx_file(path) -> Document:
    document = Document(path)
    return document

def iter_block_items(parent):
    """
    Duyệt các phần tử trong tài liệu (paragraph hoặc table) theo thứ tự xuất hiện.
    """
    for child in parent.element.body.iterchildren():
        if child.tag.endswith('}p'):
            yield Paragraph(child, parent)
        elif child.tag.endswith('}tbl'):
            yield Table(child, parent)

def extract_table_with_context_by_index(doc_path, table_index=0, context_window=2):
    document = Document(doc_path)
    blocks = list(iter_block_items(document))
    table = {}
    current_table_count = -1  # dùng để đánh dấu bảng thứ mấy

    for i, block in enumerate(blocks):
        
        if isinstance(block, Table):
            current_table_count += 1
            if current_table_count == table_index:
                # ✅ Lấy văn bản trước bảng
                context = []
                j = i - 1
                while j >= 0 and len(context) < context_window:
                    if isinstance(blocks[j], Paragraph):
                        text = blocks[j].text.strip()
                        if text:
                            context.insert(0, text)  # thêm vào đầu để đúng thứ tự
                    j -= 1

                table_text = []
                for row in block.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    table_text.append(" | ".join(cells))
                table_str = "\n".join(table_text)
                
                # print(table_str)
                
                table["title"] = "\n".join(context)
                table["content"] = table_str
                table["table_index"] = table_index
                # ✅ Trả về text + bảng
                return table
    
    return None  # nếu không tìm thấy bảng theo index

def find_similar(keyword,lines,similarity_threshold,search):
    keyword_embedding = model_embedding.encode(keyword, convert_to_tensor=True)
    matched_lines = []

    # Tạo embedding cho từng dòng và tính độ tương đồng
    for line_index in range(len(lines)):
        line_embedding = model_embedding.encode(lines[line_index], convert_to_tensor=True)
        similarity = util.cos_sim(keyword_embedding, line_embedding).item()
        
        # Nếu độ tương đồng vượt ngưỡng và dòng có chứa từ liên quan đến báo cáo
        if similarity >= similarity_threshold and re.search(search, lines[line_index], re.IGNORECASE):
            matched_lines.append((lines[line_index], similarity, line_index))

    matched_lines.sort(key=lambda x: x[1], reverse=True)

    return matched_lines[0][0] 
from docx2pdf import convert

def extract_company_and_report_name(path):
    keyword_name_report = "BÁO CÁO TÀI CHÍNH"
    keyword_name_company = "CỔ PHẦN"
    search_report = r'báo cáo tài chính|hợp nhất|quý|năm'
    search_company = r'công ty|ngân hàng|cổ phần'
    similarity_threshold = 0.2
    doc = Document(path)
    paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    # print(len(paragraphs))
    paragraphs = paragraphs[:50]
    company_name = find_similar(keyword_name_company,paragraphs,similarity_threshold,search_company)
    name_report = find_similar(keyword_name_report,paragraphs,similarity_threshold,search_report)
    print(company_name,name_report)
    return company_name,name_report

def create_json_data(path):
    name_file =  os.path.splitext(os.path.basename(path))[0]
    output_json_path = f'../Data/{name_file}.json'
    print(output_json_path)
    if os.path.exists(output_json_path) is False:
        name_company, name_report = extract_company_and_report_name(path)
        tables = [block for block in iter_block_items(Document(path)) if isinstance(block, Table)]
        chunks = []
        for i in range(len(tables)):
            chunks.append(extract_table_with_context_by_index(path,i,5))
            
        output = {
            "name_company": name_company,
            "name_report": 'c',
            "tables": [
                {
                    "title": tbl["title"],
                    "content": tbl["content"],
                    "table_index": tbl["table_index"]
                }
                for tbl in chunks
            ]
        }

        # Ghi ra file JSON
        os.makedirs(os.path.dirname(output_json_path), exist_ok=True)
        with open(output_json_path, "w", encoding="utf-8") as f:
            json.dump(output, f, ensure_ascii=False, indent=2)

        print("JSON đã được tạo xong")
        add_new_doc(output_json_path)
        print("DB đã sẵn sàng ")
        

