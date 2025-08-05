from docx import Document
import pandas as pd
from langchain.text_splitter import RecursiveCharacterTextSplitter, CharacterTextSplitter
# from langchain_community.document_loaders import PyPDFLoader, DirectoryLoader
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import GPT4AllEmbeddings
from langchain_community.document_loaders import DirectoryLoader
from langchain_community.document_loaders import UnstructuredWordDocumentLoader
# Khai bao bien
pdf_data_path = "data"
vector_db_path = "vectorstores/db_faiss"

def read_docx_file(path) -> Document:
    document = Document(path)
    return document

def extract_text_from_docx(docx_path: str) -> str:
    """
    Trích xuất tất cả văn bản từ file docx theo thứ tự paragraph tôi-tệ.
    Kết quả: 1 chuỗi text, đoạn cách nhau bởi newline.
    """
    doc = Document(docx_path)
    lines = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)

    return '\n'.join(lines)
#return the informations list of indicated table
def extract_table_info(tables, index) -> list:
    table_infos = []
    table = tables[index]
    for row in table.rows:
        row_contents = [cell.text for cell in row.cells]
        table_infos.append(row_contents)
    return table_infos

def read_excel(path):
    df = pd.read_excel(path) 
    return convert_df_to_text(df)

def extract_table(path):
    document = read_docx_file(path=path)
    tables = document.tables                #tables: A list of all tables in the file
    try:
        data_tables = []
        for index in range(len(tables)):
            table_infos = extract_table_info(tables=tables, index=index)
            df = pd.DataFrame(table_infos)
            data_tables.append(df)

    except Exception as e:
        print('error extracting single table or converting into DataFrame')
    
    create_db_from_text(data_tables)
    
def convert_df_to_text(df):
    rows = []
    for index, row in df.iterrows():
        # Chuyển từng dòng thành chuỗi văn bản rõ ràng
        row_text = " | ".join(str(cell).strip() for cell in row)
        rows.append(row_text)
    return "\n".join(rows)

def create_db_from_text(data):
    
    raw_text = ""
    for df in data:
        raw_text += convert_df_to_text(df)

    # Chia nho van ban
    text_splitter = CharacterTextSplitter(
        separator="\n",
        chunk_size=512,
        chunk_overlap=50,
        length_function=len

    )

    chunks = text_splitter.split_text(raw_text)

    # Embeding
    embedding_model = GPT4AllEmbeddings(model_file = "model\\all-MiniLM-L6-v2-f16.gguf")

    # Dua vao Faiss Vector DB
    db = FAISS.from_texts(texts=chunks, embedding=embedding_model)
    db.save_local(vector_db_path)
    print("Complete")
    return db
# text = convert_df_to_text(df)
# create_db_from_text(data_tables)
