from extract_table import get_table_dataframe
from extract_table import get_document
from extract_table import extract_table_info
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from sentence_transformers import SentenceTransformer
import os
import re
import json
import numpy as np

#checking text validation
def gibbrish_detector(text) -> bool:
    #contain special character or number --> remove
    if re.fullmatch(r"[^\w\sÀ-Ỵà-ỵ]{3,}", text):
        return True
    if len(re.findall(r"[A-Z]{2,}", text)) > 5:
        return True
    if re.search(r"[\^_~@#\$%]", text):
        return True
    return False

def is_valid_text(text) -> bool:
    if not text.strip():
        return False
    if len(text.strip()) < 5:
        return False
    if gibbrish_detector(text):
        return False
    if re.search(r"[a-zA-ZÀ-Ỵà-ỵ]", text):  # có ký tự chữ
        return True
    return False

#find the header of table
def get_table_header(document, table_index, max_lookback=5) -> str:
    #find the indicated table
    table = document.tables[table_index]
    table_element = table._element

    #scan all element in the document
    body_elements = list(document.element.body.iterchildren())

    #find the position of indicated table in XML tree
    table_position = None
    for i, el in enumerate(body_elements):
        if el == table_element:
            table_position = i
            break

    #tray back to find the header of table
    if table_position is not None:
        lookback = 0
        for i in range(table_position - 1, -1, -1):
            #only tray back in a specific range
            if lookback > max_lookback:
                break
            current_el = body_elements[i]
            #in case of this element is a paragraph
            if isinstance(current_el, CT_P):
                para_candidate = current_el.text.strip()
                #checking it text validation
                if is_valid_text(para_candidate):
                    #it might be the header of the table
                    return para_candidate
                lookback += 1

    return f"Table {table_index}"

#convert table to text for easier embedding
def table_to_text(table_data, title="") -> str:
    # table_data: List[List[str]]
    flat_lines = [", ".join(row) for row in table_data]
    table_text = f"{title}\n" + "\n".join(flat_lines)
    return table_text


#todo: create vector embedding for input file
def create_embedding_vector(output_path):
    model = SentenceTransformer('all-MiniLM-L6-v2')             #the model, change it if you want
    embeddings = []
    metadatas = []

    for idx, table in enumerate(document.tables):
        data = extract_table_info(tables=document.tables, table_index=idx)
        title = get_table_header(document=document, table_index=idx, max_lookback=5)
        text = table_to_text(table_data=data, title=title)

        vector = model.encode(text)
        embeddings.append(vector)
        metadatas.append({
            "table_index": idx,
            "title": title,
            "text": text
        })

    #save vector embedding to JSON
    save_output(base_dir, embeddings, metadatas)

def save_output(base_dir, embeddings, metadatas):
    output = []

    for vec, metadata in zip(embeddings, metadatas):
        output.append({
            'embedding': vec.tolist(),
            'metadata': metadata
        })

    embeddings_dir = os.path.abspath(os.path.join(base_dir, '..', 'embeddings'))
    os.makedirs(embeddings_dir, exist_ok=True)

    output_path = os.path.join(embeddings_dir, 'tables_embedding.json')

    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(output, f, ensure_ascii=False, indent=2)

    print(f"Embeddings saved to: {output_path}")

#read input path
try:
    base_dir = os.path.dirname(__file__)
    input_path = os.path.abspath(os.path.join(base_dir, '..', '000000014601738_VI_BaoCaoTaiChinh_KiemToan_2024_HopNhat_14032025110908.docx'))
    print('file located')

except Exception as e:
    print('cant locate file')

#read document
try:
    document = get_document(path=input_path)
    print('document reading successful')

except Exception as e:
    print('cant read document')

#embedding
try:
    create_embedding_vector(output_path=base_dir)
    print('embedding vector successful')

except Exception as e:
    print('cant create embedding vector or cant save output file')
    print(f'Error: {e}')